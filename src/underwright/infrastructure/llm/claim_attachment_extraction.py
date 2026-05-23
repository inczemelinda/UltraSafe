from __future__ import annotations

import base64
import re
from pathlib import Path
from typing import Any

from docx import Document

from underwright.application.services.claim_attachment_storage_service import (
    ClaimAttachmentStorageService,
)
from underwright.domain.claim_request import ClaimAttachmentMetadata, ClaimRequest

_IMAGE_CONTENT_TYPES = frozenset({"image/jpeg", "image/png"})
_PDF_CONTENT_TYPE = "application/pdf"
_DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_DOCUMENT_CONTENT_TYPES = frozenset({_PDF_CONTENT_TYPE, _DOCX_CONTENT_TYPE})

CLAIM_ATTACHMENT_SUMMARY_SYSTEM_PROMPT = (
    "You are an insurance claim evidence triage analyst. Produce concise, "
    "underwriter-facing evidence signals from submitted claim attachments. "
    "Base every signal only on the provided document OCR and photo observations. "
    "Do not invent facts."
)

CLAIM_ATTACHMENT_SUMMARY_USER_INSTRUCTIONS = (
    "You are given everything extracted from the submitted claim attachments: "
    "document OCR, photo observations, and extraction issues. Produce 2-6 short "
    "signals that help an underwriter understand the evidence.\n\n"
    "Claim context for comparison:\n{claim_context}\n\n"
    "Deterministic consistency checks:\n{consistency_checks}\n\n"
    "Required style:\n"
    "- First decide whether the attachment evidence supports the claim reason. "
    "Treat Claim type and Claim description as the reported incident.\n"
    "- Describe what is visibly or textually going on in the photos/documents.\n"
    "- Compare the evidence to the claim context. If the deterministic checks say "
    "a mismatch must be included, include it under Out of place / needs review.\n"
    "- For a Fire claim, photos that primarily show water/flooding/water intrusion "
    "and do not show visible fire, smoke, soot, scorching, or burn damage are not "
    "sufficient visual support for the reported fire event. Flag that clearly.\n"
    "- Flag anything out of place, inconsistent, suspicious, missing, or needing "
    "manual review.\n"
    "- Mention file names only when it helps connect a signal to evidence.\n"
    "- Do not output generic checklist categories such as Key parties, Dates, "
    "Amounts, Coverage facts, or Policy or claim references.\n"
    "- Do not write Not specified.\n"
    "- Use plain text only with this exact shape: one section label on its own "
    "line ending with a colon, followed by one dash bullet per line.\n"
    "- Prefer these exact labels when useful: Evidence signals, Out of place / "
    "needs review, Follow-up.\n"
    "- Never put bullets inline after a section label.\n"
    "- Never combine multiple bullets onto one line.\n"
    "- Omit any section that has no useful signals.\n"
    "- If all evidence is inconclusive, say that directly in one bullet.\n\n"
    "Submitted attachment evidence:\n\n{evidence}"
)


class DisabledClaimAttachmentTextExtractor:
    def __init__(self, reason: str) -> None:
        self.reason = reason

    def extract_texts(
        self,
        claim_request: ClaimRequest,
        storage: ClaimAttachmentStorageService,
    ) -> list[dict[str, Any]]:
        return [
            {
                "file_name": attachment.file_name,
                "content_type": attachment.content_type,
                "text": "",
                "error": self.reason,
            }
            for attachment in claim_request.attachments
        ]


class DisabledClaimAttachmentSummaryGenerator:
    def __init__(self, reason: str) -> None:
        self.reason = reason

    def summarize(
        self,
        extraction_results: list[dict[str, Any]],
        *,
        claim_context: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        return {"summary": "", "key_info": {}, "error": self.reason}


class OpenAIClaimAttachmentTextExtractor:
    def __init__(
        self,
        *,
        api_key: str,
        model: str = "gpt-4.1-mini",
    ) -> None:
        from openai import OpenAI

        self.client = OpenAI(api_key=api_key)
        self.model = model

    def extract_texts(
        self,
        claim_request: ClaimRequest,
        storage: ClaimAttachmentStorageService,
    ) -> list[dict[str, Any]]:
        results: list[dict[str, Any]] = []
        for attachment in claim_request.attachments:
            try:
                text = self._extract_text_from_attachment(attachment, storage)
                results.append(
                    {
                        "file_name": attachment.file_name,
                        "content_type": attachment.content_type,
                        "text": text,
                        "error": None,
                    }
                )
            except Exception as exc:  # noqa: BLE001
                results.append(
                    {
                        "file_name": attachment.file_name,
                        "content_type": attachment.content_type,
                        "text": "",
                        "error": str(exc),
                    }
                )
        return results

    def _extract_text_from_attachment(
        self,
        attachment: ClaimAttachmentMetadata,
        storage: ClaimAttachmentStorageService,
    ) -> str:
        storage_key = attachment.metadata.get("storage_key")
        if not storage_key:
            return ""

        stored = storage.get_attachment(str(storage_key))
        content_type = str(attachment.content_type or "").strip().lower()
        if content_type in _IMAGE_CONTENT_TYPES:
            return self._extract_from_image(stored.path, content_type)
        if content_type in _DOCUMENT_CONTENT_TYPES:
            return self._extract_from_document(stored.path, content_type)
        return ""

    def _extract_from_image(self, path: Path, content_type: str) -> str:
        image_base64 = base64.b64encode(path.read_bytes()).decode("utf-8")
        response = self.client.responses.create(
            model=self.model,
            input=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "input_text",
                            "text": (
                                "Analyze this insurance claim incident image for claim evidence. "
                                "Return only dash-prefixed plain-text fields using these labels "
                                "when applicable: Visible damage, Affected area, Damage severity, "
                                "Relevant conditions, Readable text, Photo uncertainty. Say "
                                "unclear when the image does not support a field. Do not identify "
                                "people, infer causes beyond visible evidence, or invent facts."
                            ),
                        },
                        {
                            "type": "input_image",
                            "image_url": f"data:{content_type};base64,{image_base64}",
                        },
                    ],
                }
            ],
        )
        return str(response.output_text or "")

    def _extract_from_document(self, path: Path, content_type: str) -> str:
        if content_type == _DOCX_CONTENT_TYPE:
            return _extract_from_docx(path)
        if content_type == _PDF_CONTENT_TYPE:
            return self._extract_from_pdf(path)
        return ""

    def _extract_from_pdf(self, path: Path) -> str:
        pdf_data = base64.b64encode(path.read_bytes()).decode("utf-8")
        response = self.client.responses.create(
            model=self.model,
            input=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "input_text",
                            "text": (
                                "Extract all visible text from this insurance claim PDF. "
                                "Do not ignore text in images, stamps, tables, or scanned regions. "
                                "Return only the extracted text."
                            ),
                        },
                        {
                            "type": "input_file",
                            "filename": path.name,
                            "file_data": f"data:application/pdf;base64,{pdf_data}",
                        },
                    ],
                }
            ],
        )
        return str(response.output_text or "")


class OpenAIClaimAttachmentSummaryGenerator:
    def __init__(
        self,
        *,
        api_key: str,
        model: str = "gpt-4o-mini",
        client: Any | None = None,
    ) -> None:
        if client is None:
            from openai import OpenAI

            client = OpenAI(api_key=api_key)
        self.client = client
        self.model = model

    def summarize(
        self,
        extraction_results: list[dict[str, Any]],
        *,
        claim_context: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        evidence_text = self._evidence_prompt_text(extraction_results)
        if not evidence_text:
            return {
                "summary": "No claim evidence signals were extracted.",
                "key_info": {},
                "error": None,
            }

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": CLAIM_ATTACHMENT_SUMMARY_SYSTEM_PROMPT,
                    },
                    {
                        "role": "user",
                        "content": CLAIM_ATTACHMENT_SUMMARY_USER_INSTRUCTIONS.format(
                            claim_context=self._claim_context_prompt_text(
                                claim_context
                            ),
                            consistency_checks=self._consistency_checks_prompt_text(
                                claim_context,
                                evidence_text,
                            ),
                            evidence=evidence_text
                        ),
                    },
                ],
                temperature=0,
            )
            summary_text = self._format_summary_text(
                response.choices[0].message.content or ""
            )
            return {
                "summary": summary_text,
                "key_info": {"raw_response": summary_text},
                "error": None,
            }
        except Exception as exc:  # noqa: BLE001
            return {"summary": "", "key_info": {}, "error": str(exc)}

    def _format_summary_text(self, raw_summary: str) -> str:
        text = str(raw_summary or "").strip()
        if not text:
            return ""

        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(
            r"(Evidence signals|Out of place\s*/\s*needs review|Out of place needs review|Follow[- ]up|Needs review|Recommendation|Recommended next step)\s*:",
            self._summary_section_replacement,
            text,
            flags=re.IGNORECASE,
        )
        text = re.sub(r":\s*[-*]\s+", ":\n- ", text)
        text = re.sub(r"([.!?])\s+[-*]\s+(?=[A-Z0-9\"'])", r"\1\n- ", text)
        text = re.sub(r"\s+[-*]\s+(?=[A-Z0-9\"'])", "\n- ", text)
        return "\n".join(line.strip() for line in text.splitlines() if line.strip())

    def _summary_section_replacement(self, match: re.Match[str]) -> str:
        prefix = "" if match.start() == 0 or match.string[match.start() - 1] == "\n" else "\n"
        return f"{prefix}{self._canonical_summary_label(match.group(1))}:"

    def _canonical_summary_label(self, value: str) -> str:
        normalized = " ".join(value.lower().replace("-", " ").replace("/", " ").split())
        labels = {
            "evidence signals": "Evidence signals",
            "follow up": "Follow-up",
            "needs review": "Needs review",
            "out of place needs review": "Out of place / needs review",
            "recommendation": "Recommendation",
            "recommended next step": "Recommended next step",
        }
        return labels.get(normalized, value.strip())

    def _claim_context_prompt_text(
        self,
        claim_context: dict[str, Any] | None,
    ) -> str:
        if not claim_context:
            return "- No claim context was provided. Do not infer a mismatch beyond the attachment evidence."

        labels = {
            "claim_type": "Claim type",
            "description": "Claim description",
            "incident_date": "Incident date",
            "incident_time": "Incident time",
            "estimated_damage": "Estimated damage",
            "policy_number": "Policy number",
            "property_address": "Property address",
        }
        lines: list[str] = []
        for key, label in labels.items():
            value = claim_context.get(key)
            if value is None or not str(value).strip():
                continue
            lines.append(f"- {label}: {value}")
        return "\n".join(lines) or "- No usable claim context was provided."

    def _consistency_checks_prompt_text(
        self,
        claim_context: dict[str, Any] | None,
        evidence_text: str,
    ) -> str:
        if not claim_context:
            return "- No claim context was provided, so no incident-type mismatch check can be made."

        claim_reason = " ".join(
            [
                str(claim_context.get("claim_type") or ""),
                str(claim_context.get("description") or ""),
            ]
        ).lower()
        evidence = evidence_text.lower()

        if self._mentions_fire(claim_reason) and self._mentions_water_damage(evidence):
            return (
                "- MUST include this mismatch under Out of place / needs review: "
                "the claim reason is Fire, but the uploaded attachment observations "
                "primarily describe water/flood damage. State that the photos do not "
                "appear related to the reported fire event unless the client explains "
                "how the water damage resulted from the fire response."
            )

        return "- No deterministic incident-type mismatch was detected from the extracted observations."

    def _mentions_fire(self, text: str) -> bool:
        return any(
            term in text
            for term in (
                "fire",
                "smoke",
                "burn",
                "burned",
                "burnt",
                "soot",
                "scorch",
                "incend",
            )
        )

    def _mentions_water_damage(self, text: str) -> bool:
        return any(
            term in text
            for term in (
                "water",
                "flood",
                "flooding",
                "water intrusion",
                "water ingress",
                "wet floor",
                "standing water",
                "leak",
                "moisture",
            )
        )

    def _evidence_prompt_text(self, extraction_results: list[dict[str, Any]]) -> str:
        blocks = [
            block
            for result in extraction_results
            if (block := self._evidence_prompt_block(result))
        ]
        return "\n\n".join(blocks)

    def _evidence_prompt_block(self, result: dict[str, Any]) -> str:
        file_name = str(result.get("file_name") or "attachment").strip()
        content_type = str(result.get("content_type") or "").strip()
        attachment_kind = self._attachment_kind(content_type)
        text = str(result.get("text") or "").strip()
        error = str(result.get("error") or "").strip()
        lines = [
            f"--- {file_name} ---",
            f"Attachment type: {attachment_kind}",
        ]
        if content_type:
            lines.append(f"Content type: {content_type}")
        if text:
            lines.append("Extracted observations:")
            lines.append(text)
        if error:
            lines.append(f"Extraction issue: {error}")
        if len(lines) <= 3:
            return ""
        return "\n".join(lines)

    def _attachment_kind(self, content_type: str) -> str:
        normalized = content_type.lower()
        if normalized in _IMAGE_CONTENT_TYPES:
            return "photo"
        if normalized in _DOCUMENT_CONTENT_TYPES:
            return "document"
        return "attachment"


def _extract_from_docx(path: Path) -> str:
    doc = Document(path)
    text_parts: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n".join(text_parts)


__all__ = [
    "CLAIM_ATTACHMENT_SUMMARY_SYSTEM_PROMPT",
    "CLAIM_ATTACHMENT_SUMMARY_USER_INSTRUCTIONS",
    "DisabledClaimAttachmentSummaryGenerator",
    "DisabledClaimAttachmentTextExtractor",
    "OpenAIClaimAttachmentSummaryGenerator",
    "OpenAIClaimAttachmentTextExtractor",
]
